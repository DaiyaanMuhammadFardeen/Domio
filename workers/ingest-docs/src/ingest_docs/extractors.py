"""Document extractors — PDF, DOCX, Notion, Markdown."""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ExtractedDocument:
    """A document after extraction, before chunking."""

    kind: str
    title: str
    paragraphs: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    metadata: dict[str, object] = field(default_factory=dict)


def extract_pdf(path: str | Path) -> ExtractedDocument:
    """Extract text from a PDF file using PyMuPDF (fitz).

    Falls back to a pytesseract OCR hook if the page has very little text.
    """
    try:
        import fitz  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError(
            "PyMuPDF is required for PDF extraction. "
            "Install with: pip install 'ingest-docs[pdf]'"
        ) from exc

    doc_path = Path(path)
    doc = fitz.open(str(doc_path))
    paragraphs: list[str] = []
    for page in doc:
        text = str(page.get_text())
        if text and text.strip():
            paragraphs.append(text.strip())
        else:
            # Short/no text → might need OCR.
            ocr_text = _try_ocr_hook(page)
            if ocr_text:
                paragraphs.append(ocr_text)
    doc.close()

    return ExtractedDocument(
        kind="pdf",
        title=doc_path.stem,
        paragraphs=paragraphs,
        metadata={"page_count": len(paragraphs)},
    )


def _try_ocr_hook(page: object) -> str:
    """Attempt OCR on a page with little extractable text.

    This is a hook point — returns empty string if pytesseract is not installed.
    """
    try:
        import pytesseract  # type: ignore[import-not-found]
        from PIL import Image

        pix = page.get_pixmap()  # type: ignore[attr-defined]
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        return pytesseract.image_to_string(img).strip()  # type: ignore[no-any-return]
    except ImportError:
        return ""
    except Exception:  # noqa: BLE001
        return ""


def extract_docx(path: str | Path) -> ExtractedDocument:
    """Extract text from a DOCX file using python-docx.

    Preserves paragraphs and tables (tables stored as rows).
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Install with: pip install 'ingest-docs[docx]'"
        ) from exc

    doc_path = Path(path)
    doc = Document(str(doc_path))
    paragraphs: list[str] = []
    tables: list[list[list[str]]] = []

    # Use python-docx high-level API for body elements.
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        table_data: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_data.append(cells)
        if table_data:
            tables.append(table_data)

    return ExtractedDocument(
        kind="docx",
        title=doc_path.stem,
        paragraphs=paragraphs,
        tables=tables,
    )


def extract_notion(blocks: list[dict[str, object]]) -> ExtractedDocument:
    """Extract text from Notion block tree.

    Supports: paragraph, heading_1..3, bulleted_list_item,
    numbered_list_item, table, code, quote.
    Walks nested children recursively.
    """
    paragraphs: list[str] = []

    def _walk(block_list: list[dict[str, object]]) -> None:
        for block in block_list:
            rich_text = block.get("rich_text", [])
            text = _flatten_rich_text(rich_text)  # type: ignore[arg-type]
            if text:
                paragraphs.append(text)

            # Recurse into children.
            children = block.get("children", [])
            if children:
                _walk(children)  # type: ignore[arg-type]

    _walk(blocks)
    return ExtractedDocument(kind="notion", title="notion-import", paragraphs=paragraphs)


def _flatten_rich_text(items: list[dict[str, object]]) -> str:
    """Concatenate Notion rich_text array into plain text."""
    parts: list[str] = []
    for item in items:
        plain = item.get("plain_text", "")
        if plain:
            parts.append(str(plain))
    return "".join(parts)


def extract_markdown(text: str) -> ExtractedDocument:
    """Extract text from Markdown, stripping code fences and YAML frontmatter."""
    # Strip YAML frontmatter.
    text = re.sub(r"^---\n.*?\n---\n", "", text, count=1, flags=re.DOTALL)
    # Strip code fences.
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    # Collapse multiple blank lines.
    text = re.sub(r"\n{3,}", "\n\n", text)

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    return ExtractedDocument(kind="markdown", title="markdown-import", paragraphs=paragraphs)
