"""Tests for the ingest-docs extractors module."""
from __future__ import annotations

from pathlib import Path

from ingest_docs.extractors import extract_docx, extract_markdown, extract_notion


def test_extract_docx(tmp_path: Path) -> None:
    """Build a tiny DOCX via python-docx and verify extraction."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        return  # skip if python-docx not installed

    doc = Document()
    doc.add_paragraph("Hello world")
    doc.add_paragraph("Second paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"

    docx_path = tmp_path / "test.docx"
    doc.save(str(docx_path))

    result = extract_docx(docx_path)
    assert result.kind == "docx"
    assert len(result.paragraphs) >= 2
    assert any("Hello world" in p for p in result.paragraphs)
    assert len(result.tables) == 1
    assert result.tables[0][0] == ["A", "B"]


def test_extract_markdown_basic() -> None:
    """Markdown extraction strips code fences and frontmatter."""
    md = """---
title: Test
---

# Heading

Some content here.

```python
code = True
```

## More

Trailing text.
"""
    result = extract_markdown(md)
    assert result.kind == "markdown"
    # Frontmatter and code fence should be stripped.
    assert not any("---" in p for p in result.paragraphs)
    assert not any("```" in p for p in result.paragraphs)
    # Content paragraphs preserved.
    texts = " ".join(result.paragraphs)
    assert "Heading" in texts
    assert "Some content" in texts
    assert "Trailing text" in texts


def test_extract_notion_nested() -> None:
    """Notion block tree with nested children is walked recursively."""
    blocks = [
        {
            "type": "heading_1",
            "rich_text": [{"plain_text": "Title"}],
            "children": [],
        },
        {
            "type": "bulleted_list_item",
            "rich_text": [{"plain_text": "Item 1"}],
            "children": [
                {
                    "type": "bulleted_list_item",
                    "rich_text": [{"plain_text": "Nested item"}],
                    "children": [],
                },
            ],
        },
        {
            "type": "paragraph",
            "rich_text": [{"plain_text": "Final paragraph"}],
            "children": [],
        },
    ]
    result = extract_notion(blocks)
    assert result.kind == "notion"
    texts = " ".join(result.paragraphs)
    assert "Title" in texts
    assert "Item 1" in texts
    assert "Nested item" in texts
    assert "Final paragraph" in texts


def test_needs_ocr_hook(tmp_path: Path) -> None:
    """Verify OCR hook is callable (returns empty string when no pytesseract)."""
    from ingest_docs.extractors import _try_ocr_hook

    # Pass a mock page object.
    result = _try_ocr_hook(object())
    # Without pytesseract installed, should return "".
    assert isinstance(result, str)
